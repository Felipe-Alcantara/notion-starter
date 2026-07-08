from __future__ import annotations

from pathlib import Path

from docx import Document

from notion_starter.services.relatorios_docx import (
    exportar_relatorios_docx,
    renderizar_docx,
)


class FakeClient:
    def __init__(self) -> None:
        self.filtro = None
        self.chamadas: list[tuple[str, object]] = []

    def consultar_database(self, database_id, *, buscar_todos=False, filtro=None):
        self.chamadas.append(("consultar_database", database_id))
        self.filtro = filtro
        assert buscar_todos is True
        return [
            {
                "id": "page1",
                "properties": {
                    "Resumo longo": {
                        "type": "rich_text",
                        "rich_text": [{"plain_text": "Texto que nao deve virar titulo"}],
                    },
                    "Nome": {"type": "title", "title": [{"plain_text": "Relatorio 06"}]},
                    "Data": {"type": "date", "date": {"start": "2026-07-06"}},
                    "Status": {"type": "status", "status": {"name": "Fechado"}},
                    "Resumo": {
                        "type": "rich_text",
                        "rich_text": [{"plain_text": "Resumo executivo"}],
                    },
                },
            }
        ]

    def ler_blocos(self, block_id, *, buscar_todos=False, recursivo=False):
        self.chamadas.append(("ler_blocos", block_id))
        assert buscar_todos is True
        assert recursivo is True
        return [
            {
                "type": "heading_1",
                "heading_1": {"rich_text": [{"plain_text": "Entrega"}]},
            },
            {
                "type": "paragraph",
                "paragraph": {"rich_text": [{"plain_text": "Texto do corpo"}]},
            },
        ]


def _texto_docx(caminho: Path) -> str:
    doc = Document(caminho)
    paragrafos = [p.text for p in doc.paragraphs]
    tabelas = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragrafos + tabelas)


def _estilos_docx(caminho: Path) -> list[str]:
    doc = Document(caminho)
    return [p.style.name for p in doc.paragraphs if p.text.strip() and p.style is not None]


def test_exporta_intervalo_lendo_propriedades_e_corpo(tmp_path: Path):
    cliente = FakeClient()

    resultado = exportar_relatorios_docx(
        database_id="db1",
        data_inicio="2026-07-01",
        data_fim="2026-07-06",
        saida=tmp_path,
        cliente=cliente,
    )

    assert resultado["total"] == 1
    assert resultado["periodo"] == {"de": "2026-07-01", "ate": "2026-07-06"}
    assert cliente.filtro == {
        "and": [
            {"property": "Data", "date": {"on_or_after": "2026-07-01"}},
            {"property": "Data", "date": {"on_or_before": "2026-07-06"}},
        ]
    }
    arquivo = Path(resultado["arquivos"][0]["arquivo"])
    assert arquivo.exists()
    texto = _texto_docx(arquivo)
    assert "Relatorio 06" in texto
    assert "Resumo executivo" in texto
    assert "Texto do corpo" in texto
    assert "Sumario" not in texto
    assert "Relatorio completo" not in texto


def test_renderizador_converte_markdown_basico_para_docx(tmp_path: Path):
    caminho = tmp_path / "relatorio.docx"

    renderizar_docx(
        caminho,
        titulo="Relatorio exemplo",
        data_relatorio="2026-07-06",
        propriedades={"Data": "2026-07-06", "Status": "Concluido", "Bloqueios": "Nenhum"},
        markdown=(
            "# Titulo\n\n"
            "## Linha do Tempo\n\n"
            "- 09:55 - item **forte**\n"
            "- 10:27 - segunda entrega\n\n"
            "| A | B |\n| --- | --- |\n| 1 | 2 |"
        ),
    )

    texto = _texto_docx(caminho)
    assert "Relatorio exemplo" in texto
    assert "Titulo" in texto
    assert "item forte" in texto
    assert "Nenhum" in texto
    assert "A" in texto and "2" in texto
    assert "09:55" in texto and "segunda entrega" in texto
    estilos = _estilos_docx(caminho)
    assert "Title" not in estilos
    assert "Heading 1" in estilos


def test_capa_usa_acentos_e_travessao_como_nos_exemplos(tmp_path: Path):
    caminho = tmp_path / "capa.docx"

    renderizar_docx(
        caminho,
        titulo="Relatorio exemplo",
        data_relatorio="2026-07-07",
        propriedades={"Data": "2026-07-07"},
        markdown="# Corpo",
    )

    texto = _texto_docx(caminho)
    assert "Relatório de Sessão" in texto
    assert "07 de julho de 2026  —  terça-feira" in texto
    assert "Relatorio de Sessao" not in texto


def test_metadados_nao_repetem_titulo_nem_secoes_de_destaque(tmp_path: Path):
    caminho = tmp_path / "metadados.docx"

    renderizar_docx(
        caminho,
        titulo="Relatorio exemplo",
        data_relatorio="2026-07-07",
        propriedades={
            "Data": "2026-07-07",
            "Relatório": "Relatorio exemplo",
            "Próximos passos": "Rotacionar o token",
            "Resumo": "Resumo executivo",
        },
        markdown="# Corpo",
    )

    doc = Document(caminho)
    celulas_chave = [row.cells[0].text for table in doc.tables for row in table.rows]
    assert "Próximos passos" not in celulas_chave
    assert "Relatório" not in celulas_chave
    texto = _texto_docx(caminho)
    # As propriedades ocultadas da tabela continuam no documento como seções.
    assert "Rotacionar o token" in texto
    assert "Resumo executivo" in texto


def test_identidade_visual_segue_os_exemplos(tmp_path: Path):
    caminho = tmp_path / "visual.docx"

    renderizar_docx(
        caminho,
        titulo="Relatorio exemplo",
        data_relatorio="2026-07-07",
        propriedades={"Data": "2026-07-07", "Status": "Concluído"},
        markdown="# Corpo\n\n| A | B |\n| --- | --- |\n| 1 | 2 |",
    )

    doc = Document(caminho)
    normal = doc.styles["Normal"]
    assert normal.font.name == "Arial"
    assert normal.font.size.pt == 10
    assert str(normal.font.color.rgb) == "404040"
    h1 = doc.styles["Heading 1"]
    assert str(h1.font.color.rgb) == "1B3A5C"
    assert h1.font.bold is True

    # Capa: projeto grande em marinho e status verde (Concluído).
    capa = [p for p in doc.paragraphs if p.text.strip()][:5]
    assert capa[0].runs[0].font.size.pt == 24
    assert str(capa[0].runs[0].font.color.rgb) == "1B3A5C"
    status = next(p for p in capa if p.text.startswith("Status:"))
    assert str(status.runs[0].font.color.rgb) == "1A5C2E"

    # Tabelas: bordas explicitas, chave da tabela de metadados sombreada e
    # cabecalho da tabela markdown em marinho com texto branco.
    metadados, markdown_tbl = doc.tables[0], doc.tables[1]
    assert "tblBorders" in metadados._tbl.xml
    assert "D5E8F0" in metadados._tbl.xml
    assert "1B3A5C" in markdown_tbl._tbl.xml
    cabecalho = markdown_tbl.rows[0].cells[0]
    assert str(cabecalho.paragraphs[0].runs[0].font.color.rgb) == "FFFFFF"


def test_periodo_invertido_falha(tmp_path: Path):
    cliente = FakeClient()

    try:
        exportar_relatorios_docx(
            database_id="db1",
            data_inicio="2026-07-06",
            data_fim="2026-07-01",
            saida=tmp_path,
            cliente=cliente,
        )
    except ValueError as exc:
        assert "data_fim" in str(exc)
    else:
        raise AssertionError("periodo invertido deveria falhar")
