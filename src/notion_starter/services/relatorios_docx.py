"""Exportacao de relatorios diarios do Notion para DOCX.

O caso de uso junta propriedades e corpo da pagina numa unica passada: consulta
o database por periodo, le os blocos de cada relatorio como Markdown e renderiza
um arquivo DOCX profissional por dia. A borda CLI/API deve apenas resolver
parametros; a regra de negocio fica aqui.
"""

from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any

try:  # pragma: no cover - coberto indiretamente nos ambientes com dependencia.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise RuntimeError(
        "A exportacao DOCX exige a dependencia 'python-docx'. "
        "Instale o pacote ou reinstale notion-starter."
    ) from exc

from notion_starter import readers
from notion_starter.content import blocos_para_markdown

# Identidade visual extraida dos DOCXs de exemplo anexados na task de origem
# (relatorios de 2026-04-22 a 2026-05-01): Arial no corpo, azul-marinho nos
# titulos, azul nos subtitulos, texto em cinza-escuro e tabelas com cabecalho
# preenchido + linhas zebradas.
_FONTE = "Arial"
_COR_MARINHO = RGBColor(0x1B, 0x3A, 0x5C)
_COR_AZUL = RGBColor(0x2E, 0x75, 0xB6)
_COR_AZUL_ESCURO = RGBColor(0x1F, 0x4D, 0x78)
_COR_TEXTO = RGBColor(0x40, 0x40, 0x40)
_COR_DATA = RGBColor(0x59, 0x59, 0x59)
_COR_VERDE = RGBColor(0x1A, 0x5C, 0x2E)
_FILL_CABECALHO = "1B3A5C"
_FILL_CHAVE = "D5E8F0"
_FILL_ZEBRA = "F2F2F2"
_FILL_CODIGO = "F2F2F2"

PROPRIEDADES_DESTAQUE_PADRAO = (
    "Data",
    "Status",
    "Area",
    "Área",
    "Resumo",
    "Bloqueios",
    "Proximos passos",
    "Próximos passos",
    "Tasks feitas",
    "Arquivos complementares",
)

SECOES_DESTAQUE = (
    ("Resumo do Dia", ("Resumo",)),
    ("Bloqueios", ("Bloqueios",)),
    ("Próximos passos", ("Proximos passos", "Próximos passos")),
)


@dataclass(frozen=True)
class RelatorioDocxExportado:
    """Resumo serializavel de um DOCX gerado."""

    id: str
    titulo: str
    data: str
    arquivo: str


def exportar_relatorios_docx(
    *,
    database_id: str,
    data_inicio: str,
    data_fim: str,
    saida: str | Path,
    cliente: Any,
    campo_data: str = "Data",
    propriedades_destaque: tuple[str, ...] = PROPRIEDADES_DESTAQUE_PADRAO,
) -> dict[str, Any]:
    """Exporta um arquivo DOCX por relatorio no intervalo informado.

    Args:
        database_id: ID do database de relatorios.
        data_inicio: Inicio do periodo em ISO ``YYYY-MM-DD``.
        data_fim: Fim do periodo em ISO ``YYYY-MM-DD``.
        saida: Diretorio onde os arquivos serao gravados.
        cliente: Instancia compativel com ``NotionClient``.
        campo_data: Nome da propriedade de data usada no filtro.
        propriedades_destaque: Propriedades que aparecem primeiro no DOCX.

    Returns:
        Dicionario serializavel para CLI/API, com periodo e arquivos gerados.
    """

    inicio = _parse_data(data_inicio, "data_inicio")
    fim = _parse_data(data_fim, "data_fim")
    if fim < inicio:
        raise ValueError("data_fim deve ser maior ou igual a data_inicio.")

    destino = Path(saida)
    destino.mkdir(parents=True, exist_ok=True)

    filtro = {
        "and": [
            {"property": campo_data, "date": {"on_or_after": inicio.isoformat()}},
            {"property": campo_data, "date": {"on_or_before": fim.isoformat()}},
        ]
    }
    paginas = cliente.consultar_database(database_id, buscar_todos=True, filtro=filtro)
    relatorios = [
        _montar_relatorio(
            pagina,
            cliente=cliente,
            destino=destino,
            campo_data=campo_data,
            propriedades_destaque=propriedades_destaque,
        )
        for pagina in paginas
    ]
    relatorios.sort(key=lambda item: (item.data, item.titulo.lower()))

    return {
        "database_id": database_id,
        "periodo": {"de": inicio.isoformat(), "ate": fim.isoformat()},
        "campo_data": campo_data,
        "saida": str(destino),
        "total": len(relatorios),
        "arquivos": [asdict(item) for item in relatorios],
    }


def _montar_relatorio(
    pagina: dict[str, Any],
    *,
    cliente: Any,
    destino: Path,
    campo_data: str,
    propriedades_destaque: tuple[str, ...],
) -> RelatorioDocxExportado:
    valores = readers.extrair_valores(pagina)
    titulo = _titulo_pagina(pagina, valores)
    data_relatorio = _data_pagina(valores, campo_data)
    blocos = cliente.ler_blocos(pagina["id"], buscar_todos=True, recursivo=True)
    markdown = blocos_para_markdown(blocos)

    nome_arquivo = f"{data_relatorio} - {_nome_seguro(titulo or 'Relatorio')}.docx"
    caminho = destino / nome_arquivo
    renderizar_docx(
        caminho,
        titulo=titulo or "Relatorio",
        data_relatorio=data_relatorio,
        propriedades=valores,
        markdown=markdown,
        propriedades_destaque=propriedades_destaque,
    )
    return RelatorioDocxExportado(
        id=pagina["id"],
        titulo=titulo,
        data=data_relatorio,
        arquivo=str(caminho),
    )


def renderizar_docx(
    caminho: str | Path,
    *,
    titulo: str,
    data_relatorio: str,
    propriedades: dict[str, Any],
    markdown: str,
    propriedades_destaque: tuple[str, ...] = PROPRIEDADES_DESTAQUE_PADRAO,
) -> Path:
    """Renderiza um relatorio ja carregado em um arquivo DOCX."""

    documento = Document()
    _configurar_documento(documento)
    _adicionar_capa(documento, titulo, data_relatorio, propriedades)
    _adicionar_tabela_metadados(
        documento, propriedades, propriedades_destaque, titulo=titulo
    )
    _adicionar_secoes_destaque(documento, propriedades)
    _adicionar_markdown(documento, markdown)

    destino = Path(caminho)
    destino.parent.mkdir(parents=True, exist_ok=True)
    documento.save(destino)
    return destino


def _configurar_documento(documento: Any) -> None:
    section = documento.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = documento.styles
    normal = styles["Normal"]
    normal.font.name = _FONTE
    normal.font.size = Pt(10)
    normal.font.color.rgb = _COR_TEXTO
    normal.paragraph_format.space_before = Pt(3)
    normal.paragraph_format.space_after = Pt(4)

    headings = (
        ("Heading 1", 16, _COR_MARINHO, 17, 5),
        ("Heading 2", 13, _COR_AZUL, 14, 4),
        ("Heading 3", 12, _COR_AZUL_ESCURO, 10, 3),
    )
    for style_name, size, cor, antes, depois in headings:
        estilo = styles[style_name]
        estilo.font.name = _FONTE
        estilo.font.size = Pt(size)
        estilo.font.bold = True
        estilo.font.color.rgb = cor
        estilo.paragraph_format.space_before = Pt(antes)
        estilo.paragraph_format.space_after = Pt(depois)

    for style_name in ("List Bullet", "List Number"):
        estilo = styles[style_name]
        estilo.font.name = _FONTE
        estilo.font.size = Pt(10)
        estilo.font.color.rgb = _COR_TEXTO
        estilo.paragraph_format.space_before = Pt(2)
        estilo.paragraph_format.space_after = Pt(2)


def _sombrear_celula(celula: Any, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    celula._tc.get_or_add_tcPr().append(shd)


def _bordas_tabela(tabela: Any) -> None:
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borda = OxmlElement(f"w:{lado}")
        borda.set(qn("w:val"), "single")
        borda.set(qn("w:sz"), "4")
        borda.set(qn("w:color"), "BFBFBF")
        borders.append(borda)
    tabela._tbl.tblPr.append(borders)


def _texto_celula(
    celula: Any, texto: str, *, bold: bool = False, cor: RGBColor | None = None
) -> None:
    p = celula.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(texto)
    run.bold = bold
    if cor is not None:
        run.font.color.rgb = cor


def _adicionar_capa(
    documento: Any,
    titulo: str,
    data_relatorio: str,
    propriedades: dict[str, Any],
) -> None:
    dia = _dia_semana(data_relatorio)
    projeto = _primeiro_valor(propriedades, ("Projeto", "Projetos", "Area"))
    status = _valor_para_texto(_primeiro_valor(propriedades, ("Status",)))
    # (texto, tamanho, cor, negrito) — hierarquia visual dos exemplos: projeto
    # grande em marinho, subtitulo azul, data em cinza e status em verde quando
    # concluido.
    linhas: list[tuple[str, int, RGBColor, bool]] = [
        (
            _valor_para_texto(projeto).upper() if projeto else "VITIS SOULS",
            24,
            _COR_MARINHO,
            True,
        ),
        ("Relatório de Sessão", 18, _COR_AZUL, False),
        (f"{_data_extenso(data_relatorio)}  —  {dia}", 13, _COR_DATA, False),
    ]
    if titulo:
        linhas.append((titulo, 12, _COR_TEXTO, True))
    if status:
        cor_status = _COR_VERDE if status.lower().startswith("conclu") else _COR_DATA
        linhas.append((f"Status: {status}", 11, cor_status, True))

    for indice, (texto, tamanho, cor, negrito) in enumerate(linhas):
        p = documento.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24 if indice == 0 else 2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(texto)
        run.bold = negrito
        run.font.size = Pt(tamanho)
        run.font.color.rgb = cor

    documento.add_paragraph()


def _adicionar_tabela_metadados(
    documento: Any,
    propriedades: dict[str, Any],
    propriedades_destaque: tuple[str, ...],
    *,
    titulo: str = "",
) -> None:
    # Fora da tabela: propriedades que já viram seção de destaque (nas duas
    # grafias) e o título, que já aparece na capa — os exemplos não os repetem.
    ocultas = {"Resumo", "Bloqueios", "Proximos passos", "Próximos passos"}
    ordenadas = [
        (nome, valor)
        for nome, valor in _ordenar_propriedades(propriedades, propriedades_destaque)
        if nome not in ocultas
        and _valor_para_texto(valor)
        and _valor_para_texto(valor) != titulo
    ]
    if not ordenadas:
        return
    tabela = documento.add_table(rows=0, cols=2)
    _bordas_tabela(tabela)
    for nome, valor in ordenadas:
        cells = tabela.add_row().cells
        _texto_celula(cells[0], nome, bold=True, cor=_COR_MARINHO)
        _sombrear_celula(cells[0], _FILL_CHAVE)
        _texto_celula(cells[1], _valor_para_texto(valor))
    documento.add_paragraph()


def _adicionar_secoes_destaque(documento: Any, propriedades: dict[str, Any]) -> None:
    for titulo, nomes in SECOES_DESTAQUE:
        valor = _primeiro_valor(propriedades, nomes)
        if valor:
            documento.add_heading(titulo, level=1)
            _adicionar_paragrafos_texto(documento, _valor_para_texto(valor))


def _adicionar_markdown(documento: Any, markdown: str) -> None:
    linhas = markdown.splitlines()
    i = 0
    em_codigo = False
    codigo: list[str] = []
    titulo_atual = ""
    while i < len(linhas):
        linha = linhas[i]
        if linha.startswith("```"):
            if em_codigo:
                _adicionar_codigo(documento, "\n".join(codigo))
                codigo = []
                em_codigo = False
            else:
                em_codigo = True
            i += 1
            continue
        if em_codigo:
            codigo.append(linha)
            i += 1
            continue
        if _linha_tabela(linha):
            bloco, i = _coletar_tabela(linhas, i)
            _adicionar_tabela_markdown(documento, bloco)
            continue
        if _linha_timeline(linha) and "linha do tempo" in titulo_atual.lower():
            bloco, i = _coletar_timeline(linhas, i)
            _adicionar_tabela_timeline(documento, bloco)
            continue
        _adicionar_linha_markdown(documento, linha)
        titulo = _titulo_markdown(linha)
        if titulo:
            titulo_atual = titulo
        i += 1
    if codigo:
        _adicionar_codigo(documento, "\n".join(codigo))


def _adicionar_linha_markdown(documento: Any, linha: str) -> None:
    texto = linha.strip()
    if not texto:
        documento.add_paragraph()
        return
    if texto.startswith("# "):
        documento.add_heading(texto[2:].strip(), level=1)
        return
    if texto.startswith("## "):
        documento.add_heading(texto[3:].strip(), level=2)
        return
    if texto.startswith("### "):
        documento.add_heading(texto[4:].strip(), level=3)
        return
    if texto.startswith("- [ ] ") or texto.startswith("- [x] "):
        marcador = "[x]" if texto.startswith("- [x] ") else "[ ]"
        p = documento.add_paragraph(style="List Bullet")
        _adicionar_inline(p, f"{marcador} {texto[6:]}")
        return
    if texto.startswith("- "):
        p = documento.add_paragraph(style="List Bullet")
        _adicionar_inline(p, texto[2:])
        return
    if re.match(r"^\d+\.\s+", texto):
        p = documento.add_paragraph(style="List Number")
        _adicionar_inline(p, re.sub(r"^\d+\.\s+", "", texto))
        return
    if texto.startswith("> "):
        p = documento.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(texto[2:]).italic = True
        return
    p = documento.add_paragraph()
    _adicionar_inline(p, texto)


def _adicionar_inline(paragrafo: Any, texto: str) -> None:
    padrao = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in padrao.finditer(texto):
        if match.start() > pos:
            paragrafo.add_run(texto[pos : match.start()])
        token = match.group(0)
        run = paragrafo.add_run(token[2:-2] if token.startswith("**") else token[1:-1])
        if token.startswith("**"):
            run.bold = True
        else:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        pos = match.end()
    if pos < len(texto):
        paragrafo.add_run(texto[pos:])


def _adicionar_codigo(documento: Any, texto: str) -> None:
    p = documento.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), _FILL_CODIGO)
    p._p.get_or_add_pPr().append(shd)


def _adicionar_paragrafos_texto(documento: Any, texto: str) -> None:
    for linha in texto.splitlines() or [texto]:
        if linha.strip():
            documento.add_paragraph(linha.strip())


def _linha_tabela(linha: str) -> bool:
    return linha.strip().startswith("|") and linha.strip().endswith("|")


def _coletar_tabela(linhas: list[str], inicio: int) -> tuple[list[str], int]:
    bloco: list[str] = []
    i = inicio
    while i < len(linhas) and _linha_tabela(linhas[i]):
        bloco.append(linhas[i])
        i += 1
    return bloco, i


def _adicionar_tabela_markdown(documento: Any, linhas: list[str]) -> None:
    linhas_validas = [
        [celula.strip() for celula in linha.strip().strip("|").split("|")]
        for linha in linhas
        if not re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", linha.strip())
    ]
    if not linhas_validas:
        return
    tabela = documento.add_table(rows=0, cols=len(linhas_validas[0]))
    _bordas_tabela(tabela)
    for numero, linha in enumerate(linhas_validas):
        cells = tabela.add_row().cells
        for idx, valor in enumerate(linha[: len(cells)]):
            if numero == 0:
                # Cabecalho como nos exemplos: fundo marinho e texto branco.
                _texto_celula(cells[idx], valor, bold=True, cor=RGBColor(0xFF, 0xFF, 0xFF))
                _sombrear_celula(cells[idx], _FILL_CABECALHO)
            else:
                _texto_celula(cells[idx], valor)
                if numero % 2 == 0:
                    _sombrear_celula(cells[idx], _FILL_ZEBRA)
    documento.add_paragraph()


def _titulo_markdown(linha: str) -> str:
    texto = linha.strip()
    if texto.startswith("# "):
        return texto[2:].strip()
    if texto.startswith("## "):
        return texto[3:].strip()
    if texto.startswith("### "):
        return texto[4:].strip()
    return ""


def _linha_timeline(linha: str) -> bool:
    texto = linha.strip()
    return bool(re.match(r"^[-*]\s+\d{1,2}:\d{2}(?:[–-]\d{1,2}:\d{2})?\s+[-–—:]\s+", texto))


def _coletar_timeline(linhas: list[str], inicio: int) -> tuple[list[tuple[str, str]], int]:
    bloco: list[tuple[str, str]] = []
    i = inicio
    while i < len(linhas) and _linha_timeline(linhas[i]):
        texto = re.sub(r"^[-*]\s+", "", linhas[i].strip())
        match = re.match(r"^(\d{1,2}:\d{2}(?:[–-]\d{1,2}:\d{2})?)\s+[-–—:]\s+(.+)$", texto)
        if match:
            bloco.append((match.group(1), match.group(2)))
        i += 1
    return bloco, i


def _adicionar_tabela_timeline(documento: Any, linhas: list[tuple[str, str]]) -> None:
    for horario, descricao in linhas:
        tabela = documento.add_table(rows=1, cols=2)
        _bordas_tabela(tabela)
        cells = tabela.rows[0].cells
        _texto_celula(cells[0], horario, bold=True, cor=_COR_MARINHO)
        _sombrear_celula(cells[0], _FILL_CHAVE)
        _texto_celula(cells[1], _texto_sem_marcacao_inline(descricao))
    documento.add_paragraph()


def _texto_sem_marcacao_inline(texto: str) -> str:
    texto = re.sub(r"\*\*([^*]+)\*\*", r"\1", texto)
    texto = re.sub(r"`([^`]+)`", r"\1", texto)
    return texto


def _parse_data(valor: str, campo: str) -> date:
    try:
        return date.fromisoformat(valor)
    except ValueError as exc:
        raise ValueError(f"{campo} deve estar no formato YYYY-MM-DD.") from exc


def _data_pagina(valores: dict[str, Any], campo_data: str) -> str:
    valor = valores.get(campo_data)
    if not valor:
        raise ValueError(f"Relatorio sem propriedade de data preenchida: {campo_data}.")
    return str(valor)[:10]


def _titulo_pagina(pagina: dict[str, Any], valores: dict[str, Any]) -> str:
    for prop in pagina.get("properties", {}).values():
        if prop.get("type") == "title":
            titulo = readers.ler_title(prop)
            if titulo:
                return titulo
    for valor in valores.values():
        if isinstance(valor, str) and valor:
            return valor
    return pagina.get("id", "")


def _primeiro_valor(propriedades: dict[str, Any], nomes: tuple[str, ...]) -> Any:
    for nome in nomes:
        valor = propriedades.get(nome)
        if valor:
            return valor
    return None


def _ordenar_propriedades(
    propriedades: dict[str, Any], destaque: tuple[str, ...]
) -> list[tuple[str, Any]]:
    vistos: set[str] = set()
    ordenadas: list[tuple[str, Any]] = []
    for nome in destaque:
        if nome in propriedades and nome not in vistos:
            ordenadas.append((nome, propriedades[nome]))
            vistos.add(nome)
    for nome in sorted(propriedades):
        if nome not in vistos:
            ordenadas.append((nome, propriedades[nome]))
    return ordenadas


def _valor_para_texto(valor: Any) -> str:
    if valor is None:
        return ""
    if isinstance(valor, list):
        return ", ".join(_valor_para_texto(item) for item in valor if item)
    if isinstance(valor, dict):
        return ", ".join(f"{k}: {_valor_para_texto(v)}" for k, v in valor.items() if v)
    return str(valor)


def _nome_seguro(valor: str) -> str:
    nome = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", valor).strip(" .-")
    nome = re.sub(r"\s+", " ", nome)
    return nome[:120] or "Relatorio"


def _dia_semana(data_iso: str) -> str:
    nomes = (
        "segunda-feira",
        "terça-feira",
        "quarta-feira",
        "quinta-feira",
        "sexta-feira",
        "sábado",
        "domingo",
    )
    return nomes[datetime.fromisoformat(data_iso[:10]).weekday()]


def _data_extenso(data_iso: str) -> str:
    meses = (
        "janeiro",
        "fevereiro",
        "março",
        "abril",
        "maio",
        "junho",
        "julho",
        "agosto",
        "setembro",
        "outubro",
        "novembro",
        "dezembro",
    )
    data = datetime.fromisoformat(data_iso[:10])
    return f"{data.day:02d} de {meses[data.month - 1]} de {data.year}"
